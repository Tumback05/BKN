
import docx
import os
import re

# See the corresponding Word doc for overall documentation

###########################
# Language specific texts #
###########################

# Blocks: </p> will be inserted when joining
san_block_m = ('''\n
               <p>\n
               Dans le cadre de l'aide à soi-même et aux camarades, il a reçu la formation de premiers secours. \n
               Dans le cadre du cours spécialisé de secouriste d'unité, il a obtenu le certificat NAEMT 
               Trauma First Responder (TFR), qui lui a été décerné. \n
               et a été formé au Tactical Combat Casualty Care (TCCC) niveau 3.\n''')

san_block_w = ('''\n
               <p>\n
               Dans le cadre de l'aide à soi-même et aux camarades, il a reçu la formation de premiers secours. \n
               Dans le cadre du cours de secouriste d'unité, elle a obtenu le certificat NAEMT 
               Trauma First Responder (TFR) \n
               et a été formé au Tactical Combat Casuality Care (TCCC) niveau 3 \n''')

dd_block_m = ('''\n
              <p>\n
              En tant que militaire en service long, il a rempli ses obligations de service d’instruction \n
              et n’est donc plus convoqué aux cours de répétition. \n''')

dd_block_w = ('''\n
              <p>\n
              En tant que militaire en service long, elle a rempli ses obligations de service d’instruction \n
              et n’est donc plus convoqué aux cours de répétition. \n''')

# keywords to find competences in word
training_text = "suivants dans le cadre"  # in KSK mal versions have "Ausbildung" instead of "Fachausbildung"
training_test_2 = "a suivi les modules de formation" # in fr mal anders beschrieben
activities_text = "Les activités suivantes faisaient partie"
activities_text_2 = "les tâches suivantes"

# keyword to find function in word
# as well as in replace_function in html
function_text = "Fonction:"


#########
# Paths #
#########

file_path = os.path.join(os.getcwd(), 'BKN_Dokumenten', 'fr')

folders = {
    '20230725_BODLUV Br 33': True,
    '20231004_G_Rttg_ABC': True,
    '20231205_LVb Inf': True,
    '20240123_LVb FU': True
}

# gets all the names of the folders with documents in it
if len(folders) == 0:
    for folder in os.listdir(file_path):
        folders[folder] = True

male_template = os.path.join('HTML', "Template_1_Spalte_letzte_Seite_m.html")
female_template = os.path.join('HTML', "Template_1_Spalte_letzte_Seite_w.html")

# specific testing variables
current = 0

path = os.path.join(file_path, list(folders.keys())[current])

sample_doc = "210820_Sdt_BKN_KSK_Gren Aufkl_m_d.docx"
sample_fail_doc = "230130_Sdt_BKN_LVbInf_Inf BesInf Pz Fahr_m_d.docx"


#############
# Functions #
#############

def replace_str(string: str) -> str:
    """replaces umlaut as well as known mistakes within the html string"""

    html_codes = {
        'Ä': '&#196;', 'ä': '&#228;',
        'Ö': '&#214;', 'ö': '&#246;',
        'Ü': '&#220;', 'ü': '&#252;',
        'à': '&#224;', 'â': '&#226;', 'æ': '&#230;', 'ç': '&#231;', 'è': '&#232;', 'é': '&#233;', 'ê': '&#234;', 'ë': '&#235;', 'î': '&#238;', 'ï': '&#239;', 'ô': '&#244;', 
        'À': '&#192;', 'Â': '&#194;', 'Æ': '&#198;', 'Ç': '&#199;', 'È': '&#201;', 'É': '&#201;', 'Ê': '&#202;', 'Ë': '&#203;', 'Î': '&#206;', 'Ï': '&#207;', 'Ô': '&#212;', 
	    'œ': '&#339;', 'ù': '&#249;', 'û': '&#251;', 'ÿ': '&#255;',
	    'Œ': '&#338;','Ù': '&#217;', 'Û': '&#219;', 'Ÿ': '&#376;',
	    '«': '&laquo;', '»': '&raquo;',  '–': '&ndash;', "'": '&apos;', '’': '&rsquo;'
    }
    for umlaut, code in html_codes.items():
        string = string.replace(umlaut, code)
    return string


def replace_competence(sdt_competences: list, html_v1: str) -> str:
    """replaces the competence description from the template html with the one from word"""

    doc_split = html_v1.split("<ul>")
    training_block = doc_split[1]
    task_block = doc_split[2]

    training_lines = training_block.split('\n') if training_block else []
    training_index = 0
    for line in training_lines:
        if "</ul>" in line:
            break
        training_index += 1
    training_lines = training_lines[training_index:]
    if not sdt_competences:
        print("Error: 'sdt_competences' is an empty list.")
        return
    new_training_block = '\n' + sdt_competences[0] + '\n'.join(training_lines)
    doc_split[1] = new_training_block

    task_lines = task_block.split('\n')
    task_index = 0
    for line in task_lines:
        if "</ul>" in line:
            break
        task_index += 1
    task_lines = task_lines[task_index:]

    new_task_block = '\n' + sdt_competences[1] + '\n'.join(task_lines)
    doc_split[2] = new_task_block

    html_v2 = "<ul>".join(doc_split)
    return html_v2


def replace_end(is_einh_san: bool, is_dd: bool, html_v2: str, is_male: bool):
    """If is DD or san, additional text will be added in the ending part of the html (see 'Language specific texts')"""

    if not is_einh_san and not is_dd:
        return html_v2
    html_blocks = html_v2.split("</ul>")
    additional_block = html_blocks[2]
    end_blocks = additional_block.split("</p>")
    if is_einh_san:
        end_blocks[0] = san_block_m if is_male else san_block_w
    if is_dd:
        end_blocks.insert(-1, dd_block_m if is_male else dd_block_w)

    additional_block = "</p>".join(end_blocks)
    html_blocks[2] = additional_block
    html_v3 = "</ul>".join(html_blocks)
    return html_v3


def create_html_file(doc_attributes: list, title: str, is_male: bool, create_file: bool, print_html: bool):
    """Gets attributes of word, creates html out of it and saves that, if 'create_file'"""

    path_to_template = path_to_male_template if is_male else path_to_female_template
    with open(path_to_template, 'r') as f:
        html_v0 = f.read()
    sdt_function = doc_attributes[0]
    html_v1 = html_v0.replace('{{User.Function}}', sdt_function)
    sdt_competences = doc_attributes[1]
    html_v2 = replace_competence(sdt_competences, html_v1)
    is_einh_san = doc_attributes[2]
    is_dd = doc_attributes[3]
    html_v3 = replace_end(is_einh_san, is_dd, html_v2, is_male)
    html_v4 = replace_str(html_v3)
    if print_html:
        print(f'Created file successfully: ', title)
    if create_file:
        with open(title, 'w') as f:
            f.write(html_v4)


def find_competence_cell(word_doc: docx, doc_name: str):
    """returns the cell of the Word with the competence description"""

    for table in word_doc.tables:
        for row in table.rows:
            for paragraph in row.cells[0].paragraphs:
                if training_text in paragraph.text:
                    return row.cells[0]
                elif training_test_2 in paragraph.text:
                    return row.cells[0]
    print(f'Error, couldnt find competence cell. Word: ', path, doc_name)
    return None


def competence_from_word(doc_name: str) -> list:
    """Collects all the description elements of the word and returns it"""

    list_space = '						'
    indented_list_space = '							'
    end_of_indented_list = list_space + '</ul>' + '\n'

    word_doc = docx.Document(os.path.join(path, doc_name))
    sdt_competence_cell = find_competence_cell(word_doc, doc_name)
    sdt_competences = ['']
    double_indent = False

    for paragraph in sdt_competence_cell.paragraphs:
        if activities_text in paragraph.text or activities_text_2 in paragraph.text:
            if double_indent:
                double_indent = False
                sdt_competences[-1] += end_of_indented_list
            sdt_competences.append('')
        elif not paragraph.text.strip():
            continue
        elif paragraph.style.name == 'List Paragraph':
            # double indent lists
            try:
                if paragraph._element.xpath('.//w:ilvl')[0].val == 1:
                    if not double_indent:
                        double_indent = True
                        sdt_competences[-1] += list_space + '<ul class="a">' + '\n'
                    sdt_competences[-1] += indented_list_space + '<li>' + paragraph.text + '</li>' + '\n'
            except:
                problematic_docs[miscellaneous].append(doc_name)
            # single indent list
            else:
                if double_indent:
                    double_indent = False
                    sdt_competences[-1] += end_of_indented_list
                sdt_competences[-1] += list_space + '<li>' + paragraph.text + '</li>' + '\n'

    if double_indent:
        sdt_competences[-1] += end_of_indented_list

    incomplete_competences = len(sdt_competences) < 2
    if incomplete_competences:
        print("ERROR: " + doc_name)
        problematic_docs[competence_txt].append(doc_name)
        return []
    
    no_competence_text = sdt_competences[0].count('\n') < 2 or sdt_competences[1].count('\n') < 2
    if no_competence_text:
        problematic_docs[competence_txt].append(doc_name)

    return sdt_competences


def function_from_word(doc_name: str) -> str:
    """Extracts the function name out of the word"""

    doc = docx.Document(os.path.join(path, doc_name))

    sdt_function = None
    for paragraph in doc.paragraphs:
        if function_text in paragraph.text:
            sdt_function = re.sub(function_text, "", paragraph.text).strip()
            break
    return sdt_function


def create_html_path(doc_name: str) -> str:
    """Properly formats the title and adds a path"""

    html_name = doc_name
    html_name = html_name.replace(' ', '_')
    html_name = html_name.replace("BKN_", '')
    html_name = html_name.replace(".docx", "_BKN.html")
    html_name = re.sub("\d+_Sdt_", "", html_name)
    html_name = re.sub("__+", "_", html_name)
    html_name = html_name.replace('?', 'ue')
    html_name = html_name.replace('ü', 'ue')
    html_name = html_name.replace('„', 'ae')
    html_name = html_name.replace('ä', 'ae')
    html_path = os.path.join(path, "HTML", "TEST", html_name)

    if not html_name.isascii():
        problematic_docs[sign_txt].append(doc_name)

    return html_path


def print_word_tables(doc_name: str):
    """Prints all the tables of a word, for better debugging"""

    doc = docx.Document(os.path.join(path, doc_name))
    print("\n",
          "##########\n",
          "# TABLES #\n",
          "##########\n")
    for table in doc.tables:
        print("### TABLE ###")
        for row in table.rows:
            for cell in row.cells:
                print("# Cell #")
                for paragraph in cell.paragraphs:
                    print(paragraph.text)
    print("###############################\n")


def make_new_html(doc_name: str, create_a_file: bool, print_html: bool = False, print_tables: bool = False):
    """Extracts information out of Word and into a new html, saves that is 'create_a_file'"""

    print("Word: Path: " + path + ' - Name: ' + doc_name)
    path_to_html = create_html_path(doc_name)
    if create_a_file:
        print("Html: " + path_to_html)
    is_male = "_m_" in doc_name

    # assert "d." in doc_name
    if print_tables:
        print_word_tables(doc_name)

    doc_attributes = [
        function_from_word(doc_name),
        competence_from_word(doc_name),
        "Einh San" in doc_name,
        "DD" in doc_name
    ]
    create_html_file(doc_attributes, path_to_html, is_male, create_file = True, print_html = False)


# In case of unaccepted changes in some docs
def accept_all_changes(doc_name: str):
    """When Word docs have unaccepted changes python can't read the text, so this has to run through them first"""

    doc = docx.Document(os.path.join(path, doc_name))

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.accepted = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.accepted = True
    doc.save(path + '/' + doc_name)


def iterate_word_docs(create_folders: bool, exception_list: list):
    """given a certain 'path', this runs through every Word and calls 'make_new_html'"""

    for doc_name in os.listdir(path):
        if doc_name.endswith('.docx') and not doc_name.startswith('~$') and doc_name not in exception_list:
            make_new_html(doc_name, create_folders)



################
# Running Code #
################


# iterate_word_docs(create_folders=True, exceptions=exceptions)
# make_new_html(sample_doc, create_a_file=False, print_html=False, print_tables=False)
# accept_all_changes(sample_fail_doc)
# iterate_folders = False
iterate_folders = True

sign_txt = "Wrong character in title (not ASCII):"
competence_txt = "No competence text:"
miscellaneous = "Other errors:"
no_paragraph = "Nothing in the Table:"

problematic_docs = {
    sign_txt: [],
    competence_txt: [],
    miscellaneous: [],
    no_paragraph: []
}

exceptions = [
    '20230530_Kader_BKN_Chance Armee_m_f.docx',
    '20230530_Kader_BKN_Chance Armee_f_f.docx'
]

if iterate_folders:
    for key, value in folders.items():
        if not value:
            continue
        path = os.path.join(file_path, key)
        path_to_male_template = os.path.join(path, male_template)
        path_to_female_template = os.path.join(path, female_template)
        iterate_word_docs(create_folders=True, exception_list=exceptions)

print("\n# Problems")
for problem, docs in problematic_docs.items():
    print(problem)
    for problem_doc in docs:
        print('"' + problem_doc + '"' + ',')
